from pathlib import Path
from datetime import date

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = Path(__file__).with_name("AI_Tutor_Technical_Textbook.docx")

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "0B2545"
GRAY = "666666"
LIGHT_GRAY = "F2F4F7"
BLUE_GRAY = "E8EEF5"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    grid = tbl.tblGrid
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        tbl.insert(0, grid)
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[idx]))
            tc_w.set(qn("w:type"), "dxa")
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)


def set_run_font(run, size=None, color=None, bold=None, italic=None, name="Calibri"):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def style_document(doc):
    sec = doc.sections[0]
    sec.top_margin = Inches(1)
    sec.bottom_margin = Inches(1)
    sec.left_margin = Inches(1)
    sec.right_margin = Inches(1)
    sec.header_distance = Inches(0.492)
    sec.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    pf = normal.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(6)
    pf.line_spacing = 1.10

    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ]:
        st = styles[name]
        st.font.name = "Calibri"
        st._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        st._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        st.font.size = Pt(size)
        st.font.bold = True
        st.font.color.rgb = RGBColor.from_string(color)
        st.paragraph_format.space_before = Pt(before)
        st.paragraph_format.space_after = Pt(after)
        st.paragraph_format.line_spacing = 1.10

    footer = sec.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("AI Tutor Technical Textbook")
    set_run_font(run, size=9, color=GRAY)


def add_title_page(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(110)
    r = p.add_run("AI Tutor")
    set_run_font(r, size=31, color=INK, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Technical Textbook and Project Documentation")
    set_run_font(r, size=16, color=DARK_BLUE, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(36)
    r = p.add_run("A full technical guide to the architecture, implementation, APIs, data model, deployment, and future work of the AI-powered learning platform.")
    set_run_font(r, size=11, color=GRAY, italic=True)

    table = doc.add_table(rows=5, cols=2)
    table.style = "Table Grid"
    set_table_geometry(table, [1900, 7460])
    rows = [
        ("Project", "AI Tutor / Papyrus learning platform"),
        ("Repository", r"D:\Grad_scripts\ai-tutor"),
        ("Prepared for", "Technical project submission"),
        ("Generated on", date.today().isoformat()),
        ("Scope", "Frontend, backend API, AI pipeline, RAG, vector database, document extraction, text generation, TTS, deployment, tests, and appendices"),
    ]
    for i, (k, v) in enumerate(rows):
        table.cell(i, 0).text = k
        table.cell(i, 1).text = v
        set_cell_shading(table.cell(i, 0), BLUE_GRAY)
        for cell in table.row_cells(i):
            for p in cell.paragraphs:
                for run in p.runs:
                    set_run_font(run, size=10.5, bold=(cell == table.cell(i, 0)))
    doc.add_page_break()


def add_heading(doc, text, level=1):
    doc.add_heading(text, level=level)


def add_para(doc, text, bold_start=None):
    p = doc.add_paragraph()
    if bold_start and text.startswith(bold_start):
        r = p.add_run(bold_start)
        set_run_font(r, bold=True)
        p.add_run(text[len(bold_start):])
    else:
        p.add_run(text)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(4)
        p.add_run(item)


def add_numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.space_after = Pt(4)
        p.add_run(item)


def add_callout(doc, label, body):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [9360])
    table.style = "Table Grid"
    cell = table.cell(0, 0)
    set_cell_shading(cell, LIGHT_GRAY)
    p = cell.paragraphs[0]
    r = p.add_run(label + ": ")
    set_run_font(r, bold=True, color=DARK_BLUE)
    p.add_run(body)


def add_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    widths = widths or [int(9360 / len(headers))] * len(headers)
    set_table_geometry(table, widths)
    for idx, head in enumerate(headers):
        cell = table.cell(0, idx)
        cell.text = head
        set_cell_shading(cell, LIGHT_GRAY)
        for run in cell.paragraphs[0].runs:
            set_run_font(run, bold=True)
    for row in rows:
        cells = table.add_row().cells
        for idx, val in enumerate(row):
            cells[idx].text = str(val)
    set_table_geometry(table, widths)
    return table


def add_code(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text)
    set_run_font(r, name="Consolas", size=9, color="333333")


def build():
    doc = Document()
    style_document(doc)
    add_title_page(doc)

    add_heading(doc, "Table of Contents", 1)
    toc = [
        "1. Executive Overview",
        "2. Educational and Technical Problem",
        "3. System Architecture",
        "4. Technology Stack",
        "5. Repository Structure",
        "6. Backend Express API",
        "7. AI Pipeline Service",
        "8. Document Extraction Service",
        "9. Text Generation Service",
        "10. Vector Database and RAG",
        "11. Text-to-Speech Service",
        "12. Frontend Application",
        "13. End-to-End Workflows",
        "14. API Reference",
        "15. Database Design",
        "16. Deployment and Operations",
        "17. Testing and Quality Assurance",
        "18. Security, Privacy, and Reliability",
        "19. Limitations and Future Work",
        "20. Appendices",
    ]
    add_bullets(doc, toc)
    doc.add_page_break()

    add_heading(doc, "1. Executive Overview", 1)
    add_para(doc, "AI Tutor is a full-stack learning platform designed to convert course material into interactive learning experiences. The project combines a React learner interface, an Express user and lesson API, and a set of Python microservices for document extraction, text generation, retrieval-augmented question answering, vector search, and audio generation.")
    add_para(doc, "The system supports a student workflow in which a learner logs in, creates or resumes a lesson, uploads or selects learning content, reviews generated summaries, starts quizzes or exams, asks questions to a tutor panel, and receives audio-oriented lesson material. The backend architecture separates business records from AI processing so the platform can evolve from classroom prototype to service-oriented deployment.")
    add_callout(doc, "Core thesis", "The project is best understood as an AI-assisted learning pipeline: input content is extracted, normalized, stored, chunked, retrieved, transformed into learning artifacts, and delivered through a protected web interface.")

    add_heading(doc, "2. Educational and Technical Problem", 1)
    add_para(doc, "Modern students often receive learning material in several formats: PDF slides, DOCX notes, PPTX lectures, images, videos, and audio. The educational problem is not only access to content, but transforming that content into study actions such as summaries, flashcards, exams, explanations, and spoken review scripts.")
    add_para(doc, "The technical problem is therefore multi-modal and stateful. The application must authenticate users, organize lessons, extract text from different file types, generate AI learning artifacts, store reusable outputs, retrieve relevant context for questions, and keep the frontend responsive.")
    add_bullets(doc, [
        "Learner requirement: quickly turn uploaded material into study resources.",
        "Instructor or evaluator requirement: see a coherent project with frontend, backend, AI services, persistence, and deployment boundaries.",
        "Engineering requirement: keep services independently testable and deployable while preserving a simple product workflow.",
    ])

    add_heading(doc, "3. System Architecture", 1)
    add_para(doc, "The architecture uses a service-oriented backend. The frontend communicates with application APIs, while the pipeline service coordinates AI-related work across specialized internal services. Docker Compose defines the main deployment topology.")
    add_table(doc, ["Layer", "Components", "Responsibility"], [
        ("Frontend", "React, Vite, React Router, Bootstrap", "Learner interface, routing, protected pages, lesson workspace, quiz and exam views."),
        ("Application API", "Node.js, Express, Neon/PostgreSQL client", "Authentication, users, lessons, user progress, AI generation records."),
        ("Pipeline API", "FastAPI pipeline_service", "Central orchestration for upload, summary, questions, flashcards, transcript, RAG ask, and audio."),
        ("AI Services", "document_service, text_service, rag_service, tts_service", "Document extraction, generation, retrieval-augmented chat, and speech synthesis."),
        ("Persistence", "PostgreSQL with pgvector", "Documents, chunks, embeddings, summaries, flashcards, MCQs, transcripts, and user-facing records."),
    ], [1500, 3000, 4860])
    add_heading(doc, "3.1 Logical Architecture Diagram", 2)
    add_code(doc, "React Frontend -> Express API -> Neon/PostgreSQL\nReact Frontend -> Pipeline Service -> Document Service -> extracted text\nPipeline Service -> Vector Database Service -> PostgreSQL/pgvector\nPipeline Service -> Text Generation Service -> summaries/questions/flashcards/transcripts\nPipeline Service -> RAG Service -> Vector Database Service + Text Generation Service\nPipeline Service -> TTS Service -> WAV audio")

    add_heading(doc, "4. Technology Stack", 1)
    add_table(doc, ["Area", "Technology", "Use in Project"], [
        ("Frontend", "React 19, Vite 7, React Router 7", "Single-page application, fast local development, protected routes."),
        ("UI", "Bootstrap, React Bootstrap, react-icons, Framer Motion", "Layout, tabs, modals, iconography, UI interactions."),
        ("Backend API", "Node.js, Express 5, JWT, bcrypt", "REST API, authentication, password hashing, protected routes."),
        ("Database access", "@neondatabase/serverless, pg, ws", "PostgreSQL/Neon connectivity from the Node API."),
        ("AI APIs", "FastAPI, Pydantic, httpx", "Microservice APIs and service-to-service calls."),
        ("Vector search", "PostgreSQL, pgvector, SQLAlchemy", "Document chunk storage and retrieval-oriented embeddings."),
        ("Document AI", "PaddleOCR, Qwen2-VL, PyMuPDF, pdfplumber, python-docx, python-pptx", "Hybrid extraction from PDFs, images, DOCX, and PPTX."),
        ("Speech", "XTTS FastAPI service", "English/Arabic speech generation from stored transcripts."),
        ("Deployment", "Docker, Docker Compose, NVIDIA GPU runtime", "Local and GPU-enabled service orchestration."),
    ], [1700, 2700, 4960])

    add_heading(doc, "5. Repository Structure", 1)
    add_para(doc, "The repository is organized into frontend and backend roots. The backend root contains both the Node API and several Python services.")
    add_table(doc, ["Path", "Purpose"], [
        ("Frontend/", "Vite React application, components, CSS, routing, authentication screens, learning workspace."),
        ("Backend/src/", "Express application, controllers, route modules, middleware, database configuration."),
        ("Backend/pipeline_service/", "FastAPI orchestration service that frontend-facing AI flows can call."),
        ("Backend/document_service/", "Hybrid document extraction microservice for PDF, DOCX, PPTX, and images."),
        ("Backend/text_services/", "Generation endpoints for explanations, summaries, questions, flashcards, and TTS scripts."),
        ("Backend/vector_database_service/", "FastAPI service backed by PostgreSQL and pgvector for documents, chunks, retrieval, and stored artifacts."),
        ("Backend/rag/", "RAG chat service with per-user, per-lesson, per-document memory."),
        ("Backend/tts_service/", "XTTS-based text-to-speech service with Arabic and English model folders."),
        ("Backend/docker-compose.yml", "Main multi-service orchestration file."),
    ], [3000, 6360])

    add_heading(doc, "6. Backend Express API", 1)
    add_para(doc, "The Express API in Backend/src/app.js mounts route modules under /api/auth, /api/users, /api/lessons, /api/user-lessons, and /api/ai-generations. It uses CORS, JSON body parsing, and a port from the PORT environment variable with 5000 as the fallback.")
    add_heading(doc, "6.1 Authentication", 2)
    add_para(doc, "Registration validates full name, email, and password, checks duplicate email addresses, hashes passwords using bcrypt, inserts a user record, and returns a seven-day JWT. Login verifies the email/password combination and returns a JWT plus minimal user profile data.")
    add_bullets(doc, [
        "Password storage uses bcrypt hashes instead of plain text.",
        "JWT payload stores userId and is consumed by protected route middleware.",
        "The controller currently includes a fallback secret, which is acceptable for local testing but should be removed for production.",
    ])
    add_heading(doc, "6.2 Lessons and User Lesson Tracking", 2)
    add_para(doc, "Lesson routes implement CRUD operations for lessons. User-lesson routes track a logged-in user's relationship to lessons, allowing the system to represent progress, status, or resumed learning sessions.")
    add_heading(doc, "6.3 AI Generation Records", 2)
    add_para(doc, "The aiGenerationController stores generated content metadata for a user and lesson. Valid generation types in the Node layer are summary, quiz, and exam. This complements the Python vector service, which stores detailed summaries, flashcards, MCQs, and transcripts by document.")

    add_heading(doc, "7. AI Pipeline Service", 1)
    add_para(doc, "The pipeline service is the central integration point for AI workflows. Its FastAPI app exposes endpoints for health checks, document ingestion, summaries, questions, flashcards, transcripts, retrieval-augmented questions, and audio generation.")
    add_table(doc, ["Endpoint", "Purpose"], [
        ("GET /health", "Checks the text, vector database, RAG, TTS, and document services."),
        ("POST /pipeline/documents/add-text", "Stores raw text as a document, chunks it, and persists chunks."),
        ("POST /pipeline/documents/upload", "Extracts uploaded file text, stores the document, and chunks it."),
        ("POST /pipeline/summary", "Returns cached summary or generates and stores one."),
        ("POST /pipeline/questions", "Returns cached MCQs or generates and stores normalized questions."),
        ("POST /pipeline/flashcards", "Returns cached flashcards or generates and stores normalized cards."),
        ("POST /pipeline/transcript", "Generates English or Arabic transcript text and stores it."),
        ("POST /pipeline/ask", "Delegates document-grounded chat to the RAG service."),
        ("POST /pipeline/audio", "Reads an existing transcript and returns generated WAV audio."),
    ], [3000, 6360])
    add_heading(doc, "7.1 Caching Pattern", 2)
    add_para(doc, "Most pipeline functions check the vector database service before generating. If a document, summary, flashcards, MCQs, or transcript already exists for the requested identity fields, the pipeline returns cached data. Otherwise it generates the missing artifact, stores it, and returns the generated result.")
    add_heading(doc, "7.2 Normalization Pattern", 2)
    add_para(doc, "The pipeline includes normalization functions for MCQs and flashcards. MCQ options are converted into a consistent A/B/C/D dictionary, answers are uppercased, and explanation points can be collapsed into explanation text. Flashcards accept multiple possible generator field names such as question/front/term and answer/back/definition.")

    add_heading(doc, "8. Document Extraction Service", 1)
    add_para(doc, "The document service is a production-oriented hybrid extraction microservice. It combines native extraction, OCR, and VLM-based visual description. It exposes file-type-specific routers for PDF, image, DOCX, and PPTX extraction.")
    add_table(doc, ["Input Type", "Primary Method", "Secondary Method", "VLM Use"], [
        ("Searchable PDF", "pdfplumber native text", "Embedded image inspection", "Figures or visual pages."),
        ("Scanned PDF", "PyMuPDF rasterization", "PaddleOCR", "Low confidence or visual-heavy pages."),
        ("DOCX", "python-docx native text", "Embedded image detection", "Large embedded images."),
        ("PPTX", "python-pptx native text", "Shape/image traversal", "Pictures and charts."),
        ("Image", "PaddleOCR", "Image preprocessing", "Photos, charts, or low OCR confidence."),
    ], [1900, 2400, 2400, 2660])
    add_para(doc, "Important environment variables include DEVICE, OCR_LANG, OCR_USE_GPU, VLM_MODEL_ID, VLM_DTYPE, VLM_USE_FLASH_ATTENTION, VISUAL_CONTENT_RATIO_THRESHOLD, OCR_CONFIDENCE_THRESHOLD, MAX_UPLOAD_SIZE_MB, and PDF_RENDER_DPI.")

    add_heading(doc, "9. Text Generation Service", 1)
    add_para(doc, "The text generation service provides AI learning transformations over long_text input. It loads environment variables at startup and exposes FastAPI endpoints for summaries, explanations, questions, flashcards, English TTS scripts, and Arabic TTS translation.")
    add_table(doc, ["Endpoint", "Output"], [
        ("POST /api/summarize", "HTML summary under summary_html."),
        ("POST /api/explain", "Deep explanation under explanation."),
        ("POST /api/questions", "Generated MCQs with quantity and difficulty controls."),
        ("POST /api/flipcards", "Generated flashcards with quantity and difficulty controls."),
        ("POST /api/tts-script", "Friendly English script for speech generation."),
        ("POST /api/translate-to-arabic-tts", "Egyptian Arabic TTS-oriented text."),
        ("GET /health", "Service health response."),
    ], [3200, 6160])

    add_heading(doc, "10. Vector Database and RAG", 1)
    add_para(doc, "The vector database service stores documents, chunks, embeddings, summaries, flashcards, MCQs, and transcripts. It uses SQLAlchemy models and pgvector-compatible embedding columns. The current README notes that the scaffold can use a deterministic mock embedder, which allows immediate project execution and can later be replaced with a production embedding provider.")
    add_table(doc, ["Table", "Main Fields", "Purpose"], [
        ("documents", "id, uid, lid, did, title, source_name, full_text, language, doc_type, metadata", "Canonical stored document text and identity."),
        ("document_chunks", "id, did, chunk_index, chunk_text, page_start, page_end, section_title", "Retrieval units derived from documents."),
        ("chunk_embeddings", "chunk_id, embedding, model_name", "Vector representation for semantic retrieval."),
        ("summaries", "did, uid, lid, summary_text, summary_type, language", "Cached summary artifacts."),
        ("flashcards", "did, uid, lid, question, answer, language", "Cached flashcard artifacts."),
        ("mcqs", "did, uid, lid, question, options, correct_answer, explanation, language", "Cached question artifacts."),
        ("transcripts", "did, uid, lid, language, transcript_text", "Speech-ready scripts and generated transcript content."),
    ], [1900, 4400, 3060])
    add_heading(doc, "10.1 RAG Service", 2)
    add_para(doc, "The RAG service receives a question plus user, lesson, and document identifiers. It retrieves relevant chunks from the vector database service, reads short-term memory for the same user/lesson/document tuple, asks the text service to generate an answer from the retrieved context, and appends the turn to memory.")
    add_bullets(doc, [
        "Memory key: user_id, lesson_id, and document_id.",
        "Configurable memory size: MEMORY_MAX_TURNS, set to 12 in Docker Compose.",
        "Endpoints: POST /rag/ask, GET /rag/memory, DELETE /rag/memory, GET /health.",
    ])

    add_heading(doc, "11. Text-to-Speech Service", 1)
    add_para(doc, "The TTS service exposes GET /health and POST /tts. It is configured in Docker Compose as an XTTS FastAPI image with Arabic and English model directories mounted into the container. The pipeline service calls it only after a transcript has already been generated and stored.")
    add_para(doc, "This separation is important: transcript generation and audio generation are different lifecycle stages. The transcript is cached in the database, while audio bytes are currently generated on demand and returned as audio/wav.")

    add_heading(doc, "12. Frontend Application", 1)
    add_para(doc, "The frontend is a Vite React single-page application. Routes are defined in App.jsx and protected by a ProtectedRoute component that checks localStorage for loggedUser. Authenticated routes include /home, /mylearning, /reminder, /lesson, /examstart, and /flashcardquiz.")
    add_table(doc, ["Component", "Role"], [
        ("Home", "Landing/home dashboard after authentication."),
        ("Header", "Navigation, dark-mode toggle state, logout menu, user initials."),
        ("Login/Register", "Authentication forms and validation UI."),
        ("Mylearning", "Creates local lesson cards and navigates to the lesson workspace."),
        ("Lesson", "Coordinates sidebar visibility, selected learning mode, uploads, and content selection."),
        ("Sidebar", "Displays uploaded files, videos, audio, and generation actions."),
        ("LessonContent", "Tabbed overview, quiz, exam, and analytics view."),
        ("UploadedFile", "File upload, preview, and download behavior."),
        ("VideoPlayer/AudioPlayer", "Learning media presentation UI."),
        ("QuizFlashcards/ExamStart", "Interactive quiz and exam experiences using generated API data."),
        ("AITutorPanel", "Floating chat panel; currently contains demo response behavior."),
    ], [2800, 6560])
    add_callout(doc, "Integration note", "Some frontend components call older local endpoints such as http://127.0.0.1:8000/summarize, /questions, and /flashcards. The newer backend pipeline exposes these capabilities under /pipeline/summary, /pipeline/questions, and /pipeline/flashcards. A production integration pass should align the frontend with the pipeline service contract.")

    add_heading(doc, "13. End-to-End Workflows", 1)
    add_heading(doc, "13.1 Registration and Login", 2)
    add_numbered(doc, [
        "The user submits full name, email, and password from the frontend.",
        "The Express auth controller validates the payload and checks duplicate email addresses.",
        "The password is salted and hashed with bcrypt.",
        "A user row is inserted into PostgreSQL or Neon.",
        "A JWT is returned and the frontend stores the logged-in state.",
        "ProtectedRoute permits navigation to learning screens.",
    ])
    add_heading(doc, "13.2 Document Upload to Study Artifacts", 2)
    add_numbered(doc, [
        "The learner uploads a supported file from the lesson workspace.",
        "The pipeline service receives file metadata, identity fields, language, and extraction mode.",
        "The document service extracts text using native parsing, OCR, and VLM description as needed.",
        "The pipeline stores the full document text in the vector database service.",
        "The pipeline chunks the text using configured chunk size and overlap.",
        "Chunks are stored for later retrieval and generation workflows.",
        "Summaries, questions, flashcards, transcripts, and audio can be generated from the stored document.",
    ])
    add_heading(doc, "13.3 Retrieval-Augmented Question Answering", 2)
    add_numbered(doc, [
        "The user asks a question about a document.",
        "The pipeline forwards the request to the RAG service.",
        "RAG retrieves top-k chunks from the vector database service.",
        "RAG adds recent memory turns for the same user, lesson, and document.",
        "The text service generates an answer grounded in retrieved context.",
        "RAG returns the answer, retrieved chunks, memory key, and updated memory.",
    ])

    add_heading(doc, "14. API Reference", 1)
    add_heading(doc, "14.1 Express API", 2)
    add_table(doc, ["Route Group", "Methods", "Description"], [
        ("/api/auth", "POST /register, POST /login", "Authentication and JWT issuance."),
        ("/api/users", "GET /profile, GET /, PUT /profile, DELETE /:id", "Protected user profile and administration operations."),
        ("/api/lessons", "GET /, GET /:id, POST /, PUT /:id, DELETE /:id", "Protected lesson CRUD."),
        ("/api/user-lessons", "GET /, GET /:lessonId, POST /:lessonId, PUT /:lessonId, DELETE /:lessonId", "Protected user progress/tracking records."),
        ("/api/ai-generations", "GET /, GET /lesson/:lessonId, GET /:id, POST /, PUT /:id, DELETE /:id", "Protected generated content records."),
    ], [2300, 3500, 3560])
    add_heading(doc, "14.2 Python Microservices", 2)
    add_table(doc, ["Service", "Representative Endpoints"], [
        ("pipeline_service", "/health, /pipeline/documents/add-text, /pipeline/documents/upload, /pipeline/summary, /pipeline/questions, /pipeline/flashcards, /pipeline/transcript, /pipeline/ask, /pipeline/audio"),
        ("document_service", "/health, /health/ready, /api/v1/pdf/extract, /api/v1/image/extract, /api/v1/docx/extract, /api/v1/pptx/extract"),
        ("text_service", "/api/summarize, /api/explain, /api/questions, /api/flipcards, /api/tts-script, /api/translate-to-arabic-tts, /health"),
        ("vector_database_service", "/health, /documents, /documents/{did}, /documents/{did}/chunks, /rag/retrieve, /documents/{did}/summary, /documents/{did}/flashcards, /documents/{did}/mcqs, /documents/{did}/transcript/en, /documents/{did}/transcript/ar"),
        ("rag_service", "/health, /rag/ask, /rag/memory"),
        ("tts_service", "/health, /tts"),
    ], [2400, 6960])

    add_heading(doc, "15. Database Design", 1)
    add_para(doc, "The project uses two persistence concepts. The Node API targets user-facing records such as users, lessons, user_lesson, and ai_generations. The vector database service owns document-centered AI records such as documents, chunks, embeddings, summaries, flashcards, MCQs, and transcripts.")
    add_heading(doc, "15.1 Identity Fields", 2)
    add_bullets(doc, [
        "uid identifies the user.",
        "lid identifies the lesson.",
        "did identifies the document.",
        "The tuple uid/lid/did is used repeatedly to isolate cached learning artifacts and RAG memory.",
    ])
    add_heading(doc, "15.2 Constraints", 2)
    add_bullets(doc, [
        "documents.did is unique and indexed.",
        "document_chunks enforces uniqueness over did and chunk_index.",
        "summaries enforces uniqueness over did, uid, lid, summary_type, and language.",
        "transcripts enforces uniqueness over did, uid, lid, and language.",
        "document chunks cascade from their parent document through the foreign key relationship.",
    ])

    add_heading(doc, "16. Deployment and Operations", 1)
    add_para(doc, "Backend/docker-compose.yml defines the deployment topology. The composition includes text_service on host port 8001, PostgreSQL/pgvector on 5437, vector_database_service on 8004, rag_service on 8006, tts_service on 8002, document_service on 8003, and pipeline_service on 8005.")
    add_table(doc, ["Service", "Host Port", "Operational Notes"], [
        ("text_service", "8001", "Uses environment file from text_services/.env."),
        ("postgres", "5437", "pgvector image, ai_tutor database, persistent postgres_data volume."),
        ("vector_database_service", "8004", "Connects to PostgreSQL through DATABASE_URL."),
        ("rag_service", "8006", "Depends on vector database and text services; stores memory data in rag_data volume."),
        ("tts_service", "8002", "Uses prebuilt xtts-fastapi image, model volume mounts, GPU access."),
        ("document_service", "8003", "GPU target with model caches and temporary data volumes."),
        ("pipeline_service", "8005", "Depends on all AI services and exposes the central AI API."),
    ], [2400, 1400, 5560])
    add_code(doc, "cd Backend\ndocker compose up --build")
    add_para(doc, "For GPU-enabled document and TTS services, the host must provide NVIDIA drivers, Docker, and NVIDIA Container Toolkit. Health endpoints should be used before accepting frontend traffic.")

    add_heading(doc, "17. Testing and Quality Assurance", 1)
    add_para(doc, "The document service includes pytest tests for health and image behavior. The frontend defines a lint script through ESLint. The project should expand tests around service contracts and critical learning workflows.")
    add_table(doc, ["Area", "Existing/Recommended Checks"], [
        ("Document service", "Run pytest tests/ -v after installing requirements."),
        ("Frontend", "Run npm run lint and npm run build inside Frontend."),
        ("Pipeline service", "Add tests for cache-hit and cache-miss behavior in summary, questions, flashcards, transcript, and audio paths."),
        ("Vector service", "Add tests for document creation, chunk replacement, retrieval filtering, and artifact upserts."),
        ("Express API", "Add route tests for auth, protected middleware, lesson CRUD, and generated content ownership."),
        ("End-to-end", "Exercise upload -> extraction -> chunking -> summary -> quiz -> transcript -> audio."),
    ], [2200, 7160])

    add_heading(doc, "18. Security, Privacy, and Reliability", 1)
    add_heading(doc, "18.1 Security", 2)
    add_bullets(doc, [
        "Passwords are hashed with bcrypt before storage.",
        "Protected Express routes use JWT middleware.",
        "The frontend currently checks localStorage for loggedUser; production should validate token presence and expiry more directly.",
        "JWT fallback secrets should be removed; JWT_SECRET must be required in production.",
        "Uploaded documents can contain sensitive educational data; extraction services should enforce file size, file type, and retention controls.",
    ])
    add_heading(doc, "18.2 Reliability", 2)
    add_bullets(doc, [
        "The pipeline health endpoint aggregates downstream service health.",
        "Docker restart policies are set to unless-stopped for major services.",
        "Document and TTS services require GPU availability; readiness probes should prevent early traffic.",
        "The project uses caching to avoid repeated generation calls and to improve user-perceived responsiveness.",
    ])
    add_heading(doc, "18.3 Privacy", 2)
    add_para(doc, "The project stores full document text, generated artifacts, transcripts, and user identifiers. A deployed version should define retention policies, encryption strategy, audit logging, and a deletion workflow that removes all user-owned documents and generated outputs.")

    add_heading(doc, "19. Limitations and Future Work", 1)
    add_table(doc, ["Current Limitation", "Recommended Improvement"], [
        ("Frontend uses some demo/static behavior.", "Connect lesson creation, AI tutor chat, summaries, flashcards, and exams to the pipeline endpoints."),
        ("Some frontend URLs point to older local endpoints.", "Introduce a central API client with environment-based base URLs."),
        ("Audio is generated on demand and not cached.", "Add audio storage endpoints and metadata in the vector or application database."),
        ("Mock/deterministic embeddings may be used in scaffold mode.", "Replace with a production embedding model and create HNSW indexes after data volume grows."),
        ("JWT fallback secret exists.", "Fail startup when JWT_SECRET is not configured in production."),
        ("Limited automated tests across services.", "Add contract tests, integration tests, and frontend route tests."),
        ("Service ownership is split between Node and Python databases.", "Document ownership boundaries and use consistent IDs across the frontend, Express API, and pipeline."),
    ], [3300, 6060])

    add_heading(doc, "20. Appendices", 1)
    add_heading(doc, "Appendix A: Main File Map", 2)
    add_table(doc, ["File", "Description"], [
        ("Frontend/src/App.jsx", "Client routing and protected route composition."),
        ("Frontend/src/Components/Authentication/ProtectedRoute.jsx", "Local login-state route guard."),
        ("Frontend/src/Components/Mylearning.jsx", "Lesson card creation and resume navigation UI."),
        ("Frontend/src/Components/Lesson.jsx", "Lesson workspace state coordinator."),
        ("Frontend/src/Components/mylearningComponents/LessonContent.jsx", "Overview, quiz, exam, analytics tab container."),
        ("Backend/src/app.js", "Express app setup and route mounting."),
        ("Backend/src/controllers/authController.js", "Registration and login logic."),
        ("Backend/src/controllers/lessonController.js", "Lesson CRUD controller."),
        ("Backend/src/controllers/aiGenerationController.js", "Generated content record controller."),
        ("Backend/pipeline_service/app/main.py", "Pipeline FastAPI endpoints."),
        ("Backend/pipeline_service/app/pipeline.py", "Orchestration functions and caching behavior."),
        ("Backend/pipeline_service/app/clients.py", "HTTP clients for downstream services."),
        ("Backend/document_service/app/main.py", "Document extraction service startup and router registration."),
        ("Backend/text_services/main.py", "Text generation FastAPI endpoints."),
        ("Backend/vector_database_service/app/main.py", "Document, chunk, retrieval, and artifact endpoints."),
        ("Backend/vector_database_service/app/models.py", "SQLAlchemy data model."),
        ("Backend/rag/app/main.py", "RAG ask and memory endpoints."),
        ("Backend/tts_service/app.py", "TTS health and speech endpoint."),
        ("Backend/docker-compose.yml", "Service orchestration."),
    ], [4200, 5160])

    add_heading(doc, "Appendix B: Submission Checklist", 2)
    add_bullets(doc, [
        "Explain the problem statement and educational value.",
        "Show the architecture diagram and identify each service.",
        "Demonstrate authentication and protected routes.",
        "Demonstrate upload or text ingestion through the pipeline service.",
        "Demonstrate summary, questions, flashcards, transcript, RAG ask, and audio paths.",
        "Show database tables and how uid, lid, and did isolate user/lesson/document content.",
        "Run health checks for pipeline, document, text, vector, RAG, and TTS services.",
        "Run available tests and record any missing coverage as future work.",
        "Discuss security, privacy, deployment, and GPU requirements.",
    ])

    add_heading(doc, "Appendix C: Glossary", 2)
    add_table(doc, ["Term", "Meaning"], [
        ("RAG", "Retrieval-augmented generation: answering with retrieved document context."),
        ("OCR", "Optical character recognition: extracting text from scanned or image-based documents."),
        ("VLM", "Vision-language model: model that can describe or reason over images."),
        ("Chunk", "A smaller piece of a document used for retrieval and generation."),
        ("Embedding", "A vector representation of text used for semantic similarity search."),
        ("Transcript", "Speech-ready text generated from source material for TTS."),
        ("MCQ", "Multiple-choice question with options, answer, and optional explanation."),
    ], [2200, 7160])

    add_heading(doc, "Conclusion", 1)
    add_para(doc, "AI Tutor demonstrates a complete technical direction for an AI-assisted learning platform. Its main strength is the separation of concerns: the React frontend handles the learner experience, Express handles user-facing records and authentication, and Python microservices handle document extraction, learning artifact generation, retrieval, and audio. With tighter frontend-to-pipeline integration, expanded tests, and production hardening, the project can grow into a robust educational system.")

    doc.save(OUT)


if __name__ == "__main__":
    build()
