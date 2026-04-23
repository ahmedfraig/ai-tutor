"""
app/services/docx_service.py

Hybrid DOCX processor:
  - python-docx: native paragraphs, tables, headings
  - VLM: embedded images (charts, diagrams, photos)
"""
from __future__ import annotations

import io
import time
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from PIL import Image

from app.core.config import get_settings
from app.core.device import resolve_device
from app.core.logging import get_logger
from app.models.schemas import (
    DocumentMetadata,
    ElementType,
    ExtractionMode,
    ExtractionResponse,
    ExtractedElement,
)
from app.services.vlm_service import get_vlm_service
from app.utils.image_utils import visual_content_ratio

logger = get_logger(__name__)
settings = get_settings()


async def process_docx(
    path: Path,
    filename: str,
    file_size: int,
    mode: ExtractionMode = ExtractionMode.AUTO,
    describe_visuals: bool = True,
    vlm_prompt: str = "Describe the content of this image in detail.",
) -> ExtractionResponse:
    t0 = time.perf_counter()
    vlm_svc = get_vlm_service()
    elements: list[ExtractedElement] = []
    device = resolve_device(settings.device)

    doc = Document(str(path))
    core_props = doc.core_properties

    # ── Native text: paragraphs ───────────────────────────────────────────────
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if not text:
            continue
        style = para.style.name.lower() if para.style else ""
        etype = ElementType.HEADER if "heading" in style else ElementType.TEXT
        elements.append(ExtractedElement(
            element_type=etype,
            content=text,
            source="native",
            metadata={"paragraph_index": i, "style": style},
        ))

    # ── Native text: tables ───────────────────────────────────────────────────
    for t_idx, table in enumerate(doc.tables):
        rows = []
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                rows.append(" | ".join(cells))
        if rows:
            elements.append(ExtractedElement(
                element_type=ElementType.TABLE,
                content="\n".join(rows),
                source="native",
                metadata={"table_index": t_idx},
            ))

    # ── Embedded images → VLM ─────────────────────────────────────────────────
    if describe_visuals and mode in (ExtractionMode.AUTO, ExtractionMode.FULL, ExtractionMode.VLM_ONLY):
        for rel in doc.part.rels.values():
            if "image" in rel.reltype:
                try:
                    img_bytes = rel.target_part.blob
                    img = Image.open(io.BytesIO(img_bytes)).convert("RGB")
                    if img.width < 80 or img.height < 80:
                        continue
                    if visual_content_ratio(img) < settings.visual_content_ratio_threshold:
                        continue
                    desc = await vlm_svc.adescribe(img, vlm_prompt)
                    if desc.strip():
                        elements.append(ExtractedElement(
                            element_type=ElementType.FIGURE,
                            content=desc,
                            source="vlm",
                        ))
                except Exception as exc:
                    logger.warning("docx_image_failed", error=str(exc))

    elapsed_ms = (time.perf_counter() - t0) * 1000
    metadata = DocumentMetadata(
        filename=filename,
        file_type="docx",
        file_size_bytes=file_size,
        author=core_props.author,
        title=core_props.title,
    )
    return ExtractionResponse.build(
        metadata=metadata,
        elements=elements,
        processing_time_ms=elapsed_ms,
        device_used=device,
    )
